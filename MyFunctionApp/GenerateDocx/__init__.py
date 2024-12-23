import os
import uuid
import logging
import json
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from datetime import datetime

def fetch_scan_data_from_cosmos(scan_data):
    return scan_data

def remove_prefix(pdf_name):
    """
    Removes the 'intermediate/' prefix from the given pdf_name if it exists.

    Parameters:
        pdf_name (str): The original PDF filename, potentially with the 'intermediate/' prefix.

    Returns:
        str: The PDF filename without the 'intermediate/' prefix.
    """
    try:
        prefix = "intermediate/"
        if pdf_name.startswith(prefix):
            return pdf_name[len(prefix):]
        return pdf_name
    except Exception as e:
        # Log the exception if logging is set up, or handle it as needed
        # For simplicity, we'll just return the original pdf_name
        return pdf_name

def generate_docx_from_knowledge_scan(scan_data):
    doc = Document()
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # Add header to the first page
    header = section.first_page_header

    # Set the header margins to minimal
    section.top_margin = Inches(0.5)
    section.header_distance = Inches(0.3)

    # Path to the header image
    header_image_path = 'capture.png'  # Replace with your image path

    if os.path.exists(header_image_path):
        paragraph = header.paragraphs[0]
        run = paragraph.add_run()
        # Set width to span close to the entire page width (standard 8.5 inches - left & right margins)
        run.add_picture(header_image_path, width=Inches(7.5))  # Adjust width based on your content
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    else:
        logging.warning(f"Header image not found at {header_image_path}. Skipping header image.")

    # Add content to the DOCX file
    # Move the "Knowledge Scan" heading to the center with black bold font size 14
    heading_paragraph = doc.add_paragraph()
    heading_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    heading_run = heading_paragraph.add_run('Knowledge Scan')
    heading_run.bold = True
    heading_run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
    heading_run.font.size = Pt(14)

    # Add today's date formatted as "Month Day, Year" with bold and smaller font
    today = datetime.today().strftime('%B %d, %Y')
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_run = date_paragraph.add_run(today)
    date_run.bold = True
    date_run.font.size = Pt(10)  # Smaller font size
    
    # General notes section
    general_notes = scan_data.get('general_notes', 'No general notes provided.')
    general_notes_paragraph = doc.add_paragraph(general_notes)
    general_notes_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Optional: Set alignment if needed

    # Keywords and themes section
    
    keywords = scan_data.get('keywords', "").split(", ")
    keywords_paragraph = doc.add_heading('Keywords and themes:', level=3)
    if keywords_paragraph.runs:
        keywords_run = keywords_paragraph.runs[0]
        keywords_run.font.bold = True
        keywords_run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
        keywords_run.font.size = Pt(14)

    if keywords:
        for keyword in keywords:
            doc.add_paragraph(keyword, style='List Bullet')
    else:
        doc.add_paragraph('No keywords provided.')

    # Resources searched section
    resources_searched = scan_data.get('combined_summaries', [])
    resources_paragraph = doc.add_heading('Resources searched:', level=3)
    
    # Apply styling to the "Resources searched" heading
    if resources_paragraph.runs:
        resources_run = resources_paragraph.runs[0]
        resources_run.font.bold = True
        resources_run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
        resources_run.font.size = Pt(14)

    if resources_searched:
        for resource in resources_searched:
            pdf_name = resource['pdf_name']
            updated_pdf_name = remove_prefix(pdf_name)
            doc.add_paragraph(updated_pdf_name, style='List Bullet')
    else:
        doc.add_paragraph('No resources provided.')

    # Overall Summary section
    overall_summary = scan_data.get('overall_summary', '')
    overall_paragraph = doc.add_heading('Overall Summary:', level=2)
    
    # Apply styling to the "Overall Summary" heading
    if overall_paragraph.runs:
        overall_run = overall_paragraph.runs[0]
        overall_run.font.bold = True
        overall_run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
        overall_run.font.size = Pt(14)

    if overall_summary.strip():
        doc.add_paragraph(overall_summary.strip())  # Add the concatenated summaries
    else:
        doc.add_paragraph('No overall summary provided.')

    # Summaries section
    combined_summaries = scan_data.get('combined_summaries', [])
    summaries_paragraph = doc.add_heading('Sources Summaries:', level=2)
    
    # Apply styling to the "Sources Summaries" heading
    if summaries_paragraph.runs:
        summaries_run = summaries_paragraph.runs[0]
        summaries_run.font.bold = True
        summaries_run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
        summaries_run.font.size = Pt(14)

    if combined_summaries:
        for i, summary_info in enumerate(combined_summaries, 1):
            new_pdf_name = remove_prefix(summary_info['pdf_name'])
            doc.add_heading(f"Document {i}: {new_pdf_name}", level=3)
            doc.add_paragraph(summary_info.get('summary', 'No summary available.'))
            overall_summary += summary_info.get('summary', '') + " "  # Concatenate summaries
    else:
        doc.add_paragraph('No summaries available.')

    # Save the document to a file in the temp directory
    file_name = f"knowledge_scan_{uuid.uuid4()}.docx"
    doc_path = os.path.join('/tmp', file_name)
    doc.save(doc_path)

    return doc_path, file_name

def main(inputDocument: func.DocumentList, outputDocument: func.Out[func.Document], outputBlob: func.Out[str]):
    logging.info("GenerateDocx function triggered.")

    for doc in inputDocument:
        scan_data = fetch_scan_data_from_cosmos(doc)

        # Generate DOCX from the scan data
        doc_path, file_name = generate_docx_from_knowledge_scan(scan_data)

        # Read the content of the DOCX file
        with open(doc_path, 'rb') as file:
            doc_content = file.read()

        # Create a fixed filename for the outputBlob binding
        fixed_blob_path = f"curated/{file_name}"

        # Save the document content to the Blob using a fixed path
        outputBlob.set(doc_content)

        # Prepare the output document item for Cosmos DB
        doc_item = func.Document.from_dict({
            "id": str(uuid.uuid4()),
            "file_name": file_name,
            "blob_location": fixed_blob_path  # Store the Blob storage location in Cosmos DB
        })

        # Output to Cosmos DB docxContainer
        outputDocument.set(doc_item)

    logging.info("DOCX file generated, uploaded to Blob storage, and location saved to Cosmos DB successfully.")